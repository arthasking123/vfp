# optimized_text_window.py
from PyQt5.QtWidgets import QMainWindow, QTextEdit, QPushButton, QVBoxLayout, QWidget, QFileDialog, QDesktopWidget
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QTextCursor
import docx
from io import BytesIO
from PIL import Image
import base64
from typing import Tuple

class OptimizedTextWindow(QMainWindow):
    def __init__(self, optimized_content: str):
        super().__init__()
        self.init_ui(optimized_content)

    def init_ui(self, optimized_content: str) -> None:
        self.textEdit = QTextEdit(self)
        self.textEdit.setHtml(optimized_content)

        exportButton = QPushButton("导出到Word", self)
        exportButton.clicked.connect(self.export_to_word)

        layout = QVBoxLayout()
        layout.addWidget(self.textEdit)
        layout.addWidget(exportButton)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        # 获取屏幕尺寸
        screen = QDesktopWidget().screenNumber(QDesktopWidget().cursor().pos())
        screen_size = QDesktopWidget().screenGeometry(screen)

        # 计算窗口尺寸
        window_width = screen_size.width() // 2
        window_height = int(window_width * 3 / 4)  # 4:3 比例

        # 设置窗口尺寸和位置
        self.setGeometry(
            (screen_size.width() - window_width) // 2,
            (screen_size.height() - window_height) // 2,
            window_width,
            window_height
        )

        self.setWindowTitle('优化后的文本')

    def export_to_word(self) -> None:
        # 打开文件保存对话框
        file_path, _ = QFileDialog.getSaveFileName(self, "保存Word文档", "", "Word文档 (*.docx)")
        
        if not file_path:
            return  # 用户取消了保存操作

        doc = docx.Document()
        text_doc = self.textEdit.document()
        block = text_doc.begin()
        page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

        while block.isValid():
            cursor = QTextCursor(block)
            if cursor.charFormat().isImageFormat():
                image_format = cursor.charFormat().toImageFormat()
                base64_str = image_format.name()
                if base64_str.startswith('data:image/png;base64,'):
                    base64_str = base64_str.split(',')[1]

                image_data = base64.b64decode(base64_str)
                temp_stream = BytesIO(image_data)
                temp_image = Image.open(temp_stream)
                temp_width, temp_height = temp_image.size
                scaling_factor = page_width / temp_width
                adjusted_width = page_width
                adjusted_height = int(temp_height * scaling_factor)
                doc.add_picture(BytesIO(image_data), width=(adjusted_width), height=(adjusted_height))
            else:
                block_text = block.text()
                if block_text.strip():
                    doc.add_paragraph(block_text)

            block = block.next()

        doc.save(file_path)  # 使用用户选择的文件路径保存文档
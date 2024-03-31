import base64
import io
import os
import re
import subprocess
import sys
import tempfile
import threading
from datetime import datetime, timedelta
from io import BytesIO

from PIL import Image  # Make sure to import Image from Pillow
import docx
from docx.shared import Inches
import fitz
import ffmpeg
import openai
from openai import OpenAI
import vlc
from bs4 import BeautifulSoup
from PyQt5.QtCore import (QByteArray, QBuffer, QIODevice, QMimeData, QSize, QTimer, Qt, pyqtSignal, 
                          QThread, QMetaObject, Q_ARG, pyqtSlot)
from PyQt5.QtGui import (QColor, QClipboard, QImage, QIcon, QPainter, QPixmap, QTextCharFormat, QTextCursor, QBrush, QFont)
from PyQt5.QtWidgets import (QApplication, QDialog, QHBoxLayout, QLabel, QLineEdit, QListWidget, 
                             QListWidgetItem, QMainWindow, QMessageBox, QPushButton, QProgressBar, QSlider, 
                             QTextEdit, QVBoxLayout, QWidget, QProgressDialog, QFileDialog)
import whisper


class RichTextEditor(QTextEdit):
    def __init__(self, parent=None):
        super().__init__(parent)

    def copyRichText(self):
        cursor = self.textCursor()
        if not cursor.hasSelection():
            return

        mimeData = QMimeData()
        mimeData.setHtml(cursor.selection().toHtml())
        QApplication.clipboard().setMimeData(mimeData)

class SubtitleThread(QThread):
    finishedSignal = pyqtSignal(str)  # 信号，传递生成的 SRT 文件路径

    def __init__(self, videoPath, parent=None):
        super().__init__(parent)
        self.videoPath = videoPath

    def extractAudio(self):
        audioPath = "temp_audio.mp3"
        command = f"ffmpeg -i \"{self.videoPath}\"  -q:a 0 -map a -y \"{audioPath}\""
        os.system(command)        
        return audioPath

    def format_as_srt(self, transcription_segments):
        srt_output = ""
        for i, segment in enumerate(transcription_segments, 1):
            start = self.format_time(segment['start'])
            end = self.format_time(segment['end'])
            text = segment['text']
            srt_output += f"{i}\n{start} --> {end}\n{text}\n\n"
        return srt_output

    def format_time(self, seconds):
        """将秒转换为SRT时间格式（HH:MM:SS,MS）"""
        ms = int((seconds % 1) * 1000)
        s = int(seconds)
        h, s = divmod(s, 3600)
        m, s = divmod(s, 60)
        return f"{h:02}:{m:02}:{s:02},{ms:03}"

    def run(self):
        print('extractAudio...')
        audioPath = self.extractAudio()
        print('download model...')
        model = whisper.load_model("small")
        print('transcribing...')
        result = model.transcribe(audioPath, language="zh", task="transcribe")
        print(result)
        srtPath = "subtitles.srt"
        
        srt = self.format_as_srt(result['segments'])

        with open(srtPath, "w", encoding="utf-8") as file:
            file.write(srt)
        
        self.finishedSignal.emit(srtPath)  # 发射信号，附带 SRT 路径


class OptimizedTextWindow(QMainWindow):
    def __init__(self, optimized_content):
        super().__init__()
        self.initUI(optimized_content)

    def initUI(self, optimized_content):
        self.textEdit = QTextEdit(self)
        self.textEdit.setHtml(optimized_content)

        exportButton = QPushButton("导出到Word", self)
        exportButton.clicked.connect(self.exportToWord)

        layout = QVBoxLayout()
        layout.addWidget(self.textEdit)
        layout.addWidget(exportButton)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    def exportToWord(self):
        
        doc = docx.Document()
        text_doc = self.textEdit.document()
        block = text_doc.begin()
         # Calculate the available width for the image considering margins
        page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

        while block.isValid():
            cursor = QTextCursor(block)
            if cursor.charFormat().isImageFormat():
                image_format = cursor.charFormat().toImageFormat()
                base64_str = image_format.name()
                if base64_str.startswith('data:image/png;base64,'):
                    base64_str = base64_str.split(',')[1]


                # Decode the base64 string
                image_data = base64.b64decode(base64_str)

                temp_stream = BytesIO(image_data)
                temp_image = Image.open(temp_stream)
                temp_width, temp_height = temp_image.size

                # Calculate the width scaling factor
                scaling_factor = page_width / temp_width

                # Calculate the adjusted width and height
                adjusted_width = page_width
                adjusted_height = int(temp_height * scaling_factor)

                # 添加图片到 Word 文档
                doc.add_picture(BytesIO(image_data), width=(adjusted_width), height=(adjusted_height))

            else:
                block_text = block.text()
                if block_text.strip():  # 只有当文本块非空时才添加
                    doc.add_paragraph(block_text)  # 添加文本

            block = block.next()

        # 保存文档
        doc.save('exported_document.docx')

class VideoPlayer(QMainWindow):
    progressDialogClosed = pyqtSignal(list)

    def __init__(self):
        super().__init__()
        self.instance = vlc.Instance()
        self.player = self.instance.media_player_new()
        self.ppt_images = []
        self.subtitleThread = ''
        self.subtitleData = []
        self.initUI()
        self.last_search_text = None
        self.progressDialogClosed.connect(self.onProgressDialogClosed)


    def initUI(self):
        self.setWindowTitle('Video Formalization Processor')
        self.setGeometry(100, 100, 1200, 800)

        # 主布局分为左右两部分
        mainLayout = QHBoxLayout()

        # 左侧布局
        leftLayout = QVBoxLayout()
        
        self.loadVideoButton = QPushButton('Load Presentation Video', self)
        self.loadVideoButton.clicked.connect(self.loadVideo)
        leftLayout.addWidget(self.loadVideoButton)

        self.videoFrame = QWidget(self)
        self.videoFrame.setMinimumHeight(300)
        self.player.set_hwnd(self.videoFrame.winId())
        leftLayout.addWidget(self.videoFrame)

        self.slider = QSlider(self)
        self.slider.setOrientation(Qt.Horizontal)
        self.slider.sliderPressed.connect(self.sliderPressed)
        self.slider.sliderReleased.connect(self.sliderReleased)
        self.slider.sliderMoved.connect(self.setPosition)
        leftLayout.addWidget(self.slider)
        self.isSliderBeingDragged = False

         # 定时器更新进度条
        self.timer = QTimer(self)
        self.timer.setInterval(100)  # 每100毫秒更新一次
        self.timer.timeout.connect(self.updateSlider)
        self.timer.start()

        # 播放控制按钮
        self.playButton = QPushButton('Play', self)
        self.playButton.clicked.connect(self.player.play)

        self.pauseButton = QPushButton('Pause', self)
        self.pauseButton.clicked.connect(self.player.pause)

        self.stopButton = QPushButton('Stop', self)
        self.stopButton.clicked.connect(self.player.stop)

        # 按钮布局
        controlLayout = QHBoxLayout()
        controlLayout.addWidget(self.playButton)
        controlLayout.addWidget(self.pauseButton)
        controlLayout.addWidget(self.stopButton)
        leftLayout.addLayout(controlLayout)

        # 初始化定时器
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.updateSubtitles)
        self.timer.start(100)  # 每100毫秒更新一次

        

        self.loadPDFButton = QPushButton('Load PDF of Presentation', self)
        self.loadPDFButton.clicked.connect(self.loadPDF)
        leftLayout.addWidget(self.loadPDFButton)

        self.pdfListWidget = QListWidget(self)
        self.pdfListWidget.setVisible(False)
        self.pdfListWidget.itemDoubleClicked.connect(self.insertImage)

        leftLayout.addWidget(self.pdfListWidget)

        self.progressBar = QProgressBar(self)
        leftLayout.addWidget(self.progressBar)

        self.generateSubtitlesButton = QPushButton('Generate Subtitles', self)
        self.generateSubtitlesButton.clicked.connect(self.generateSubtitles)
        leftLayout.addWidget(self.generateSubtitlesButton)
        
        # 右侧布局
        rightLayout = QVBoxLayout()
        self.textEdit = RichTextEditor(self)       
        self.textEdit.setAcceptRichText(True)  # 启用富文本格式
        rightLayout.addWidget(self.textEdit)

        self.loadHTMLButton = QPushButton('Load HTML Document', self)  # 加载字幕的按钮
        self.loadHTMLButton.clicked.connect(self.loadHTML)

        self.saveHTMLButton = QPushButton('Save HTML Document', self)  # 加载字幕的按钮
        self.saveHTMLButton.clicked.connect(self.saveHTML)

        self.optimizeButton = QPushButton('Formalization', self)
        self.optimizeButton.clicked.connect(self.optimizeText)

        self.apiKeyInput = QLineEdit(self)
        self.apiKeyInput.setPlaceholderText("Enter OpenAI API Key")        

        rightLayout.addWidget(self.loadHTMLButton)
        rightLayout.addWidget(self.saveHTMLButton)
        rightLayout.addWidget(self.optimizeButton)
        rightLayout.addWidget(self.apiKeyInput)

        # 组合布局
        mainLayout.addLayout(leftLayout, 1)
        mainLayout.addLayout(rightLayout, 2)

        centralWidget = QWidget()
        centralWidget.setLayout(mainLayout)
        self.setCentralWidget(centralWidget)

    def optimizeText(self):
        openai.api_key = self.apiKeyInput.text()
        if openai.api_key == '':
            QMessageBox.information(self, "Info", "apikey is None")
            return
        segments = self.extractTextAndImages()
        # print(segments)
        
        # 初始化进度对话框
        self.progressDialog = QProgressDialog("Optimizing text...", "Cancel", 0, 100, self)
        self.progressDialog.setWindowModality(Qt.WindowModal)
        self.progressDialog.show()

        # 创建并启动后台线程
        thread = threading.Thread(target=self.optimizeInBackground, args=(segments,))
        thread.start()

        

    def optimizeInBackground(self, segments):
        optimized_segments = []
        client = OpenAI(api_key = openai.api_key)
        for i, segment in enumerate(segments):
            # 这里是一个示例，你的实际逻辑可能会有所不同
            # 假设 optimizeTextWithOpenAI 返回优化后的内容和进度
            if segment['type'] == 'text':
                # 移除字幕的索引标号和起止时间
                text_to_optimize = re.sub(r'\d+\n\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3}\n', '', segment['content'])

                # 构建消息
                system_message = "将以下口语化表述整理成连贯性的技术文章段落。"
                user_message = text_to_optimize

                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": user_message}
                    ]
                )

                optimized_text = response.choices[0].message.content
                optimized_segments.append(optimized_text.strip())
            elif segment['type'] == 'image':
                optimized_segments.append(segment['content'])  # 直接添加图片


            progress = (int)(100*i/len(segments))

            # 更新进度对话框
            QMetaObject.invokeMethod(self.progressDialog, "setValue", Qt.QueuedConnection, Q_ARG(int, progress))

        # 完成后关闭进度对话框
        QMetaObject.invokeMethod(self.progressDialog, "close", Qt.QueuedConnection)
        self.progressDialogClosed.emit(optimized_segments)

    @pyqtSlot(list)
    def onProgressDialogClosed(self, optimized_segments):
        self.optimizedWindow = OptimizedTextWindow("".join(optimized_segments))
        self.optimizedWindow.show()


    def optimizeTextWithOpenAI(self, text_segments):
        optimized_segments = []
        for segment in text_segments:
            if segment['type'] == 'text':
                # 移除字幕的索引标号和起止时间
                text_to_optimize = re.sub(r'\d+\n\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3}\n', '', segment['content'])

                # 构建消息
                system_message = "将以下口语化表述整理成连贯性书面语言，要求如下：1.逻辑结构清晰 2.遇到列举性说明，段落前要加标号 3.不要扩写句子，不要改变文本原意，仅把口语话的表述书面化 4.使用中文语言 5.遇到问句不要回答问题 6.返回HTML格式，每段文字外有<p>标签"
                user_message = text_to_optimize

                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": user_message}
                    ]
                )

                optimized_text = response['choices'][0]['message']['content']
                optimized_segments.append(optimized_text.strip())
            elif segment['type'] == 'image':
                optimized_segments.append(segment['content'])  # 直接添加图片
        return optimized_segments

    def extractTextAndImages(self):
        html_content = self.textEdit.toHtml()
        soup = BeautifulSoup(html_content, 'html.parser')

        segments = []
        current_text = ''  # 用于累积当前段落的文本

        for p_element in soup.find_all('p'):
            # 检查 <p> 标签内是否有 <span>，如果有则提取文本
            span_elements = p_element.find_all('span')
            for span in span_elements:
                current_text += span.get_text() + '\n'

            # 检查 <p> 标签内是否有 <img>，如果有则提取图片
            img_elements = p_element.find_all('img')
            for img in img_elements:
                # 遇到图片时，先添加累积的文本（如果有）
                if current_text:
                    segments.append({'type': 'text', 'content': current_text.strip()})
                    current_text = ''  # 重置文本累积

                # 然后添加图片的 HTML
                image_html = str(img)  # 或使用 img.prettify()
                segments.append({'type': 'image', 'content': f"<p>{image_html}</p>"})

         # 不要忘记添加最后一个累积的文本段落（如果有）
        if current_text:
            segments.append({'type': 'text', 'content': current_text.strip()})
        return segments

    def sliderPressed(self):
        self.isSliderBeingDragged = True

    def sliderReleased(self):
        newPosition = self.slider.value() / self.slider.maximum()
        self.player.set_position(newPosition)  # 设置新的播放位置
        self.isSliderBeingDragged = False        

    def setPosition(self, position):
        # 设置视频的播放位置
        self.player.set_position(position / 1000.0)

    def updateSlider(self):
        if not self.isSliderBeingDragged:
            # 更新进度条的值
            max_value = self.slider.maximum()  # 获取进度条的最大值
            current_position = self.player.get_position()
            progress_value = int(current_position * max_value)
            self.slider.setValue(progress_value)

    def loadSrt(self):
        filename, _ = QFileDialog.getOpenFileName(self, "Open SRT File", "", "SRT Files (*.srt)")
        if filename:
            self.loadSRT(filename)

    def loadHTML(self):
        filename, _ = QFileDialog.getOpenFileName(self, "Open HTML File", "", "HTML Files (*.HTML)")
        if filename:
            with open(filename, "r", encoding="utf-8") as file:
                self.textEdit.setHtml(file.read())

    def saveHTML(self):
        filename, _ = QFileDialog.getSaveFileName(self, "SAVE HTML File", "", "HTML Files (*.HTML)")
        if filename:
            with open(filename, "w", encoding="utf-8") as file:
                file.write(self.textEdit.toHtml())

    
    def loadVideo(self):
        videoPath, _ = QFileDialog.getOpenFileName(self, "Open Video", "", "Video Files (*.mp4)")
        if videoPath:
            self.videoPath = videoPath
            self.playVideo(videoPath)

    def playVideo(self, videoPath):
        media = self.instance.media_new(videoPath)
        self.player.set_media(media)
        self.player.play()

    def updateWhisperProgress(self, result, total):
        self.progressBar.setValue(int((result / total) * 100))

    def generateSubtitles(self):
        if hasattr(self, 'videoPath'):
            self.generateSubtitlesButton.setEnabled(False)
            self.subtitleThread = SubtitleThread(self.videoPath)
            self.subtitleThread.finishedSignal.connect(self.onThreadFinished)  # 连接信号
            self.subtitleThread.start()

    def onThreadFinished(self, srtPath):
        self.progressBar.setValue(0)
        self.generateSubtitlesButton.setEnabled(True)
        self.loadSRT(srtPath)
        self.playVideo(self.videoPath)

    def generateSubtitlesThread(self):
        print('extractAudio...')
        audioPath = self.extractAudio()
        print('download model...')
        model = whisper.load_model("small")
        print('transcribing...')
        result = model.transcribe(audioPath, language="zh", task="transcribe")
        print(result)
        srtPath = "subtitles.srt"
        
        srt = self.format_as_srt(result['segments'])

        with open(srtPath, "w", encoding="utf-8") as file:
            file.write(srt)

        
    def insertImage(self, item):
        listWidget = item.listWidget()
        row = listWidget.row(item)
        pixmap = self.pdf_images[row]

        cursor = self.textEdit.textCursor()
        image = pixmap.toImage()
        byte_array = QByteArray()
        buffer = QBuffer(byte_array)
        buffer.open(QIODevice.WriteOnly)
        image.save(buffer, "PNG")  # You can choose a different format if needed
        base64_data = byte_array.toBase64().data()

        cursor.insertImage(pixmap.toImage(), "data:image/png;base64,{}".format(str(base64_data, encoding='utf-8')))  # 以PNG格式插入图片
        self.textEdit.setTextCursor(cursor)

    def loadSRT(self, srtPath):
        with open(srtPath, "r", encoding="utf-8") as file:
            lines = file.readlines()

        formatted_lines = []
        previous_end_time = None

        for line in lines:
            if "-->" in line:
                start_time_str, end_time_str = line.split("-->")
                start_time = datetime.strptime(start_time_str.strip(), '%H:%M:%S,%f')
                end_time = datetime.strptime(end_time_str.strip(), '%H:%M:%S,%f')


                previous_end_time = end_time

            formatted_lines.append(line)

        self.textEdit.setText("".join(formatted_lines))

    def clearHighlight(self):
        # 清除所有高亮（恢复默认格式）
        cursor = QTextCursor(self.textEdit.document())
        cursor.select(QTextCursor.Document)
        fmt = QTextCharFormat()
        fmt.setBackground(QBrush(QColor("white")))
        cursor.mergeCharFormat(fmt)

    def highlightText(self, text):
        # 如果新的文本与上次搜索的不同，则清除之前的高亮
        if text != self.last_search_text:
            self.clearHighlight()

        # 设置高亮格式
        fmt = QTextCharFormat()
        fmt.setBackground(QBrush(QColor("yellow")))

        # 开始搜索和高亮
        cursor = QTextCursor(self.textEdit.document())
        while True:
            cursor = self.textEdit.document().find(text, cursor)
            if cursor.isNull():
                break
            cursor.mergeCharFormat(fmt)

        # 保存最后搜索的文本
        self.last_search_text = text

    def updateSubtitles(self):
        current_time = self.player.get_time()  # 获取当前视频时间，单位毫秒
        # 对应行置黄
        self.yellowLine(current_time)
        

    def yellowLine(self, current_time):
        text = self.textEdit.toPlainText()  # 从 QTextEdit 获取所有文本
        # 解析 SRT 字幕
        for block in text.split('\n\n'):
            lines = block.split('\n')
            if len(lines) >= 3:
                time_range = lines[1]
                start_time, end_time = self.parseSrtTimeRange(time_range)
                if start_time <= current_time <= end_time:
                    #对应block 置黄
                    self.highlightText('\n'.join(lines[2:]))

    def parseSrtTimeRange(self, time_range):
        # 解析 SRT 时间格式 "00:00:20,000 --> 00:00:24,400"
        match = re.match(r'(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})', time_range)
        if match:
            start_time = self.convertSrtTimeToMilliseconds(match.group(1))
            end_time = self.convertSrtTimeToMilliseconds(match.group(2))
            return start_time, end_time
        return 0, 0

    def convertSrtTimeToMilliseconds(self, srt_time):
        hours, minutes, seconds = map(int, srt_time[:-4].split(':'))
        milliseconds = int(srt_time[-3:])
        return (hours * 3600 + minutes * 60 + seconds) * 1000 + milliseconds

    def loadPDF(self):
        pdfPath, _ = QFileDialog.getOpenFileName(self, "Open PDF", "", "PDF Files (*.pdf)")
        if pdfPath:
            self.displayPDF(pdfPath)

    def displayPDF(self, pdfPath):
        doc = fitz.open(pdfPath)
        self.pdf_images = []
        self.pdfListWidget.clear()
        self.pdfListWidget.setVisible(True)

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap()
            img = QPixmap.fromImage(QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format_RGB888))
            self.pdf_images.append(img)

            # 缩放图像为原始尺寸的一半
            scaled_img = img.scaled(pix.width // 2, pix.height // 2, Qt.KeepAspectRatio, Qt.SmoothTransformation)

            item = QListWidgetItem()
            item.setSizeHint(scaled_img.size())

            label = QLabel()
            label.setPixmap(scaled_img)
            self.pdfListWidget.addItem(item)
            self.pdfListWidget.setItemWidget(item, label)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = VideoPlayer()
    ex.show()
    sys.exit(app.exec_())
